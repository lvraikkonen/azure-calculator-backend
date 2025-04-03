from typing import Dict, Any
from pathlib import Path
from app.rag.utils.file_processors.base import FileProcessorRegistry

from app.core.logging import get_logger

logger = get_logger(__name__)


@FileProcessorRegistry.register
class DocxFileProcessor:
    """Word文档处理器"""

    @classmethod
    def can_process(cls, file_path: Path) -> bool:
        suffix = file_path.suffix.lower()
        return suffix in [".doc", ".docx"]

    async def process(self, file_path: Path, encoding: str = 'utf-8', **kwargs) -> Dict[str, Any]:
        text = ""
        metadata = {}

        try:
            # 使用python-docx处理.docx文件
            if file_path.suffix.lower() == '.docx':
                import docx

                doc = docx.Document(str(file_path))

                # 提取标题
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title

                # 提取作者
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author

                # 提取创建日期
                if doc.core_properties.created:
                    metadata["created"] = str(doc.core_properties.created)

                # 提取正文
                text_parts = []
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip():
                        text_parts.append(para.text)

                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)

                text = "\n".join(text_parts)

                # 记录段落数和表格数
                metadata["paragraph_count"] = len(doc.paragraphs)
                metadata["table_count"] = len(doc.tables)

            # 对于.doc文件，尝试使用其他库或提示未实现
            else:
                logger.warning(f"不支持直接读取.doc格式，请转换为.docx: {file_path}")
                text = f"[Word文档(.doc): {file_path.name}，不支持直接读取此格式]"
                metadata["doc_support"] = False

        except ImportError:
            logger.warning("未安装python-docx，无法提取Word文档文本")
            text = f"[Word文档: {file_path.name}]"
            metadata["docx_support"] = False
        except Exception as e:
            logger.error(f"Word文档处理错误: {str(e)}")
            text = f"[Word文档处理错误: {file_path.name}]"
            metadata["processing_error"] = str(e)

        return {
            "content": text,
            "metadata": metadata
        }


@FileProcessorRegistry.register
class ExcelFileProcessor:
    """Excel文件处理器"""

    @classmethod
    def can_process(cls, file_path: Path) -> bool:
        suffix = file_path.suffix.lower()
        return suffix in [".xlsx", ".xls"]

    async def process(self, file_path: Path, encoding: str = 'utf-8', **kwargs) -> Dict[str, Any]:
        text = ""
        metadata = {}

        try:
            # 使用pandas读取Excel
            import pandas as pd

            # 读取所有sheet
            excel_file = pd.ExcelFile(str(file_path))
            sheet_names = excel_file.sheet_names
            metadata["sheet_count"] = len(sheet_names)
            metadata["sheets"] = sheet_names

            all_texts = []

            # 处理每个sheet
            for sheet_name in sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)

                # 添加sheet标题
                all_texts.append(f"\n--- Sheet: {sheet_name} ---\n")

                # 将数据框转换为字符串表示
                sheet_text = df.to_string(index=False)
                all_texts.append(sheet_text)

                # 记录行列数
                metadata[f"sheet_{sheet_name}_rows"] = len(df)
                metadata[f"sheet_{sheet_name}_columns"] = len(df.columns)

            text = "\n\n".join(all_texts)

        except ImportError:
            logger.warning("未安装pandas，无法提取Excel文件内容")
            text = f"[Excel文件: {file_path.name}]"
            metadata["excel_support"] = False
        except Exception as e:
            logger.error(f"Excel文件处理错误: {str(e)}")
            text = f"[Excel文件处理错误: {file_path.name}]"
            metadata["processing_error"] = str(e)

        return {
            "content": text,
            "metadata": metadata
        }