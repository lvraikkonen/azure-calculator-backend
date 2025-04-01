"""
文档加载器组件 - 从不同来源加载文档
"""
from typing import List, Dict, Any, Optional, Union
import aiohttp
import asyncio
import os
from pathlib import Path
import json
from datetime import datetime
import re

from app.rag.core.registry import register_component, RAGComponentRegistry
from app.rag.core.interfaces import DocumentLoader
from app.rag.core.models import Document, Metadata
from app.core.logging import get_logger

logger = get_logger(__name__)

@register_component(RAGComponentRegistry.DOCUMENT_LOADER, "web")
class WebDocumentLoader(DocumentLoader[Document]):
    """网页文档加载器 - 从URL加载网页内容"""
    
    def __init__(self, html_to_text: bool = True, timeout: int = 30, headers: Optional[Dict[str, str]] = None):
        """
        初始化网页文档加载器
        
        Args:
            html_to_text: 是否将HTML转换为纯文本
            timeout: 请求超时时间(秒)
            headers: 请求头
        """
        self.html_to_text = html_to_text
        self.timeout = timeout
        self.headers = headers or {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
    
    async def load(self, source: str, **kwargs) -> List[Document]:
        """
        加载网页文档
        
        Args:
            source: 网页URL
            **kwargs: 附加参数，可包含:
                - metadata: 自定义元数据
                - encoding: 文本编码
                - include_images: 是否包含图片文本说明
                
        Returns:
            List[Document]: 文档列表，通常只包含一个文档
        """
        pass
    
    def _html_to_text(self, html: str, include_images: bool = False) -> str:
        """
        将HTML转换为纯文本
        
        Args:
            html: HTML内容
            include_images: 是否包含图片文本说明
            
        Returns:
            str: 纯文本内容
        """
        pass
    
    def _extract_title(self, html: str) -> str:
        """
        从HTML中提取标题
        
        Args:
            html: HTML内容
            
        Returns:
            str: 标题，如果没有找到则返回空字符串
        """
        pass

@register_component(RAGComponentRegistry.DOCUMENT_LOADER, "file")
class FileDocumentLoader(DocumentLoader[Document]):
    """文件文档加载器 - 从文件系统加载文档"""
    
    def __init__(self, base_dir: Optional[str] = None):
        """
        初始化文件文档加载器
        
        Args:
            base_dir: 基础目录，如果指定，则所有文件路径都将相对于此目录
        """
        self.base_dir = Path(base_dir) if base_dir else None
    
    async def load(self, source: str, **kwargs) -> List[Document]:
        """
        加载文件文档
        
        Args:
            source: 文件路径，可以是相对路径或绝对路径
            **kwargs: 附加参数，可包含:
                - metadata: 自定义元数据
                - encoding: 文本编码
                
        Returns:
            List[Document]: 文档列表，通常只包含一个文档
        """
        # 确定文件路径
        if self.base_dir:
            file_path = self.base_dir / source
        else:
            file_path = Path(source)
        
        logger.info(f"从文件加载文档: {file_path}")
        
        # 检查文件是否存在
        if not file_path.exists():
            logger.error(f"文件不存在: {file_path}")
            return []
        
        try:
            # 获取自定义参数
            custom_metadata = kwargs.get("metadata", {})
            encoding = kwargs.get("encoding", "utf-8")
            
            # 获取文件类型
            file_type = self._get_file_type(file_path)
            
            # 根据文件类型处理文件
            if file_type == "text":
                # 处理文本文件
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                
            elif file_type == "pdf":
                # 处理PDF文件
                content = self._extract_pdf_text(file_path)
                
            elif file_type == "docx":
                # 处理Word文档
                content = self._extract_docx_text(file_path)
                
            elif file_type == "json":
                # 处理JSON文件
                with open(file_path, "r", encoding=encoding) as f:
                    data = json.load(f)
                content = json.dumps(data, ensure_ascii=False, indent=2)
                
            else:
                # 未知文件类型
                logger.warning(f"不支持的文件类型: {file_path.suffix}")
                return []
            
            # 创建元数据
            metadata = Metadata(
                source=str(file_path),
                title=file_path.name,
                created_at=datetime.fromtimestamp(file_path.stat().st_mtime),
                content_type=self._get_content_type(file_path),
                extra={
                    **custom_metadata,
                    "file_size": file_path.stat().st_size,
                    "file_type": file_type,
                    "absolute_path": str(file_path.absolute())
                }
            )
            
            # 创建文档
            document = Document(
                content=content,
                metadata=metadata
            )
            
            logger.info(f"成功加载文档: {file_path}, 内容长度: {len(content)}")
            return [document]
            
        except Exception as e:
            logger.error(f"加载文件失败: {file_path}, 错误: {str(e)}")
            return []
    
    def _get_file_type(self, file_path: Path) -> str:
        """
        获取文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 文件类型
        """
        suffix = file_path.suffix.lower()
        
        if suffix in [".txt", ".md", ".py", ".js", ".html", ".css", ".csv"]:
            return "text"
        elif suffix == ".pdf":
            return "pdf"
        elif suffix in [".doc", ".docx"]:
            return "docx"
        elif suffix == ".json":
            return "json"
        else:
            return "unknown"
    
    def _get_content_type(self, file_path: Path) -> str:
        """
        获取内容类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 内容类型
        """
        suffix = file_path.suffix.lower()
        
        if suffix == ".txt":
            return "text/plain"
        elif suffix == ".html":
            return "text/html"
        elif suffix == ".md":
            return "text/markdown"
        elif suffix == ".pdf":
            return "application/pdf"
        elif suffix in [".doc", ".docx"]:
            return "application/msword"
        elif suffix == ".json":
            return "application/json"
        elif suffix == ".csv":
            return "text/csv"
        else:
            return "application/octet-stream"
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """
        提取PDF文本
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            str: PDF文本内容
        """
        try:
            # 尝试使用PyPDF2
            from PyPDF2 import PdfReader
            
            text_parts = []
            reader = PdfReader(file_path)
            
            for page in reader.pages:
                text_parts.append(page.extract_text())
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.warning("未安装PyPDF2，无法提取PDF文本")
            return f"[PDF文件: {file_path.name}]"
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """
        提取Word文档文本
        
        Args:
            file_path: Word文档路径
            
        Returns:
            str: Word文档文本内容
        """
        try:
            # 尝试使用python-docx
            import docx
            
            doc = docx.Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                text_parts.append(para.text)
            
            return "\n".join(text_parts)
            
        except ImportError:
            logger.warning("未安装python-docx，无法提取Word文档文本")
            return f"[Word文档: {file_path.name}]"

@register_component(RAGComponentRegistry.DOCUMENT_LOADER, "azure_api")
class AzureAPIDocumentLoader(DocumentLoader[Document]):
    """Azure API文档加载器 - 从Azure API获取服务信息"""
    
    def __init__(self, api_key: Optional[str] = None, api_base: Optional[str] = None):
        """
        初始化Azure API文档加载器
        
        Args:
            api_key: Azure API密钥
            api_base: Azure API基础URL
        """
        self.api_key = api_key
        self.api_base = api_base or "https://management.azure.com"
        self.headers = {
            "User-Agent": "AzureCalculator/1.0",
            "Content-Type": "application/json"
        }
        
        if api_key:
            self.headers["Authorization"] = f"Bearer {api_key}"
    
    async def load(self, source: str, **kwargs) -> List[Document]:
        """
        从Azure API加载文档
        
        Args:
            source: API路径或服务标识符
            **kwargs: 附加参数，可包含:
                - api_version: API版本
                - metadata: 自定义元数据
                - params: 查询参数
                
        Returns:
            List[Document]: 文档列表
        """
        logger.info(f"从Azure API加载文档: {source}")
        
        pass
    
    def _create_document_from_item(
        self, 
        item: Dict[str, Any], 
        source: str, 
        item_id: str,
        custom_metadata: Dict[str, Any]
    ) -> Document:
        """
        从API响应项创建文档
        
        Args:
            item: API响应项
            source: 数据源
            item_id: 项ID
            custom_metadata: 自定义元数据
            
        Returns:
            Document: 文档
        """
        pass

@register_component(RAGComponentRegistry.DOCUMENT_LOADER, "multi")
class MultiDocumentLoader(DocumentLoader[Document]):
    """多源文档加载器 - 从多个源加载文档并合并结果"""
    
    def __init__(self, loaders: Optional[Dict[str, DocumentLoader[Document]]] = None):
        """
        初始化多源文档加载器
        
        Args:
            loaders: 加载器字典，键为类型，值为加载器实例
        """
        self.loaders = loaders or {}
    
    def add_loader(self, loader_type: str, loader: DocumentLoader[Document]):
        """
        添加加载器
        
        Args:
            loader_type: 加载器类型
            loader: 加载器实例
        """
        self.loaders[loader_type] = loader
    
    async def load(self, source: Union[str, Dict[str, Any]], **kwargs) -> List[Document]:
        """
        从多个源加载文档
        
        Args:
            source: 源描述，可以是字符串或字典
            **kwargs: 附加参数
            
        Returns:
            List[Document]: 文档列表
        """
        logger.info(f"从多个源加载文档")
        
        pass
    
    def _detect_source_type(self, source: str) -> str:
        """
        检测源类型
        
        Args:
            source: 源字符串
            
        Returns:
            str: 加载器类型
        """
        if source.startswith(("http://", "https://")):
            return "web"
        elif os.path.exists(source):
            return "file"
        elif "azure.com" in source or source.startswith(("subscriptions/", "providers/Microsoft")):
            return "azure_api"
        else:
            return "web"  # 默认为网页